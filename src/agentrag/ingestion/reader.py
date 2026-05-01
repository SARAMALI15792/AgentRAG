"""File reader — converts local files to RawDocument via the plugin registry."""

from __future__ import annotations

import hashlib
import json
import logging
from pathlib import Path
from typing import Any

import pymupdf
from bs4 import BeautifulSoup
from docx import Document

from agentrag.ingestion import reader_registry
from agentrag.types import RawDocument

logger = logging.getLogger(__name__)


def read_file(path: Path) -> RawDocument:
    """Read a local file and return its content as a RawDocument."""
    if not path.exists():
        raise FileNotFoundError(f"File not found: {path}")

    suffix = path.suffix.lower()
    reader_fn = reader_registry.get_reader(suffix)
    source_id = hashlib.sha256(str(path.resolve()).encode()).hexdigest()[:16]
    text = reader_fn(path)

    if not text.strip():
        raise ValueError(f"File is empty: {path}")

    return RawDocument(
        source_id=source_id,
        filename=path.name,
        text=text,
        metadata={},
    )


def _read_pdf(path: Path) -> str:
    """Extract all text from a PDF file using pymupdf."""
    doc = pymupdf.open(str(path))  # type: ignore[no-untyped-call]
    pages = [page.get_text() for page in doc]  # type: ignore[attr-defined]
    doc.close()  # type: ignore[no-untyped-call]
    return "\n".join(pages)


def _read_docx(path: Path) -> str:
    """Extract paragraph text from a Word document, skipping blank paragraphs."""
    doc = Document(str(path))
    lines: list[str] = [str(p.text) for p in doc.paragraphs if p.text.strip()]
    return "\n".join(lines)


def _read_html(path: Path) -> str:
    """Extract body text from HTML, removing nav/header/footer/script/style."""
    html = path.read_text(encoding="utf-8")
    soup = BeautifulSoup(html, "html.parser")
    for tag in soup.find_all(["nav", "header", "footer", "script", "style"]):
        tag.decompose()
    return soup.get_text(separator="\n", strip=True)


def _read_ipynb(path: Path) -> str:
    """Extract source text from code and markdown cells; skip raw cells."""
    # json.loads returns dict[str, Any]; notebook cell structure is heterogeneous
    nb: dict[str, Any] = json.loads(path.read_text(encoding="utf-8"))
    parts: list[str] = []
    for cell in nb.get("cells", []):
        if cell.get("cell_type") in {"code", "markdown"}:
            source = cell.get("source", [])
            # nbformat allows source as list[str] or plain str
            text = source if isinstance(source, str) else "".join(source)
            parts.append(text)
    return "\n\n".join(parts)


def _read_plaintext(path: Path) -> str:
    """Read a plain text file as UTF-8."""
    return path.read_text(encoding="utf-8")


# Register all Tier 1 readers
reader_registry.register([".txt", ".md", ".py"], _read_plaintext)
reader_registry.register([".pdf"], _read_pdf)
reader_registry.register([".docx"], _read_docx)
reader_registry.register([".html"], _read_html)
reader_registry.register([".ipynb"], _read_ipynb)
